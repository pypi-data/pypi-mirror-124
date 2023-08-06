import argparse
import os

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE

from .parse import read


def generate(in_file: str, out_file: str) -> None:
    """Read the input specification from the in_file file, and generate a Word document at out_file."""
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    def add_style(name: str, indent):
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal_style
        style.paragraph_format.left_indent = Inches(indent)
        return style

    page_style = add_style("PagePara", 0)  # noqa(F841)
    panel_style = add_style("PanelPara", 0.5)
    quote_style = add_style("QuotePara", 1.0)

    with open(in_file, "rt") as f:
        for page, panels in read(f):
            document.add_paragraph(f"{page}", style=page_style.name)
            for panel in range(panels):
                document.add_paragraph(f"PANEL {panel+1}", style=panel_style.name)
                document.add_paragraph("", style=quote_style.name)
    document.save(out_file)


def main():
    parser = argparse.ArgumentParser(
        description="Generate a Word Document from page/panel specification"
    )
    parser.add_argument("--in-file", help="input file name", required=True)
    parser.add_argument("--out-file", help="output file name")
    args = parser.parse_args()
    if args.out_file:
        out_file = args.out_file
    else:
        root, _ = os.path.splitext(args.in_file)
        out_file = root + ".docx"
    generate(args.in_file, out_file)
